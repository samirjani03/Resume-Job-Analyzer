import fitz  # PyMuPDF
import docx
import io
import re

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text clean and ordered from PDF using PyMuPDF."""
        text_content = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text_content.append(page.get_text())
        full_text = "\n".join(text_content)
        return DocumentParser._clean_text(full_text)

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX document using python-docx."""
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return DocumentParser._clean_text("\n".join(paragraphs))

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from plain text file."""
        return DocumentParser._clean_text(file_bytes.decode("utf-8", errors="ignore"))

    @staticmethod
    def parse_document(file_name: str, file_bytes: bytes) -> str:
        """Dispatcher function based on file extension."""
        ext = file_name.lower().split(".")[-1]
        if ext == "pdf":
            return DocumentParser.extract_text_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return DocumentParser.extract_text_from_docx(file_bytes)
        elif ext in ["txt", "md"]:
            return DocumentParser.extract_text_from_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _clean_text(text: str) -> str:
        """Normalize line breaks and clean excessive whitespace."""
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]{2,}', ' ', text)
        return text.strip()
